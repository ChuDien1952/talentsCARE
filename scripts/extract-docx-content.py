#!/usr/bin/env python3
"""
Extract content from Word document
"""

import sys
import json
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

def extract_docx_content(file_path):
    """Extract text content from .docx file"""
    try:
        doc = Document(file_path)

        content = {
            "paragraphs": [],
            "tables": [],
            "sections": {}
        }

        current_section = None
        section_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect section headers (usually bold or heading styles)
            if para.style.name.startswith('Heading'):
                if current_section and section_content:
                    content["sections"][current_section] = section_content
                current_section = text
                section_content = []
            else:
                if current_section:
                    section_content.append(text)
                content["paragraphs"].append({
                    "text": text,
                    "style": para.style.name
                })

        # Add last section
        if current_section and section_content:
            content["sections"][current_section] = section_content

        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            content["tables"].append(table_data)

        return content

    except Exception as e:
        print(f"ERROR: {str(e)}", file=sys.stderr)
        return None

if __name__ == "__main__":
    doc_path = Path("Documents/talentsCARE_Complete_v3_DE.docx")

    if not doc_path.exists():
        print(f"ERROR: File not found: {doc_path}", file=sys.stderr)
        sys.exit(1)

    content = extract_docx_content(doc_path)

    if content:
        # Save to JSON file
        output_path = Path("Documents/extracted_content.json")
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(content, f, ensure_ascii=False, indent=2)

        print(f"[OK] Content extracted successfully to {output_path}")
        print(f"[OK] Found {len(content['paragraphs'])} paragraphs")
        print(f"[OK] Found {len(content['sections'])} sections")
        print(f"[OK] Found {len(content['tables'])} tables")

        # Print first few sections
        print("\nSections found:")
        for i, section in enumerate(list(content['sections'].keys())[:10]):
            print(f"  {i+1}. {section}")
    else:
        sys.exit(1)
